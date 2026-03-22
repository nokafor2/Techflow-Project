#!/usr/bin/env python3
"""
Convert project DOCUMENTATION.md to Word (.docx) using python-docx.
Run from repo root: python scripts/markdown_to_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_code_font(run) -> None:
    run.font.name = "Consolas"
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    except Exception:
        pass
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_inline_paragraph(doc, text: str) -> None:
    """Add a paragraph with **bold** and `code` segments."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    # Split by **bold** and `code` — keep order
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_code_font(run)
        else:
            p.add_run(part)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", s))


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for p in table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                run = p.add_run("\n".join(code_lines))
                set_code_font(run)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        if stripped == "---":
            doc.add_paragraph("—" * 40).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Markdown table: header | --- |
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2  # skip separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            inner = stripped[2:]
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", inner)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1])
                    set_code_font(r)
                else:
                    p.add_run(part)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            rest = m.group(2)
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", rest)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1])
                    set_code_font(r)
                else:
                    p.add_run(part)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            rest = stripped[2:]
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", rest)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1])
                    set_code_font(r)
                else:
                    p.add_run(part)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Diagram / preformatted block (lines with box drawing)
        if "┌" in stripped or stripped.startswith("Developer"):
            block = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and (
                any(c in lines[i] for c in "│┌┐└┘├┤┬┴▼─") or lines[i].strip().startswith("│")
            ):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            set_code_font(run)
            continue

        # Project tree block (starts with Techflow-Project/)
        if stripped.startswith("Techflow-Project/"):
            block = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and (
                lines[i].startswith("│") or lines[i].startswith("├──") or lines[i].startswith("└──") or lines[i].strip().startswith("Techflow")
            ):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            set_code_font(run)
            continue

        # Italic last line
        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:-1])
            r.italic = True
            i += 1
            continue

        add_inline_paragraph(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


def main() -> int:
    root = Path(__file__).resolve().parent.parent
    md = root / "DOCUMENTATION.md"
    out = root / "DOCUMENTATION.docx"
    if not md.is_file():
        print(f"Not found: {md}", file=sys.stderr)
        return 1
    convert(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
